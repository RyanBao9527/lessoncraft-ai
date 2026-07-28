"""Render a concise, print-friendly lesson navigation document."""

from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from models.blueprint import CourseBlueprint
from models.lesson_plan import LessonPlan
from models.slide_deck import SlideDeck

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(33, 38, 45)
MUTED = RGBColor(93, 102, 112)
TABLE_FILL = "E8EEF5"
ALT_FILL = "F7F9FC"
LATIN_FONT = "Arial"
CJK_FONT = "STHeiti"
CODE_FONT = "Menlo"


def _set_run_font(
    run: object,
    name: str = LATIN_FONT,
    size: float | None = None,
    bold: bool | None = None,
    color: RGBColor | None = None,
    east_asia: str = CJK_FONT,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _set_cell_fill(cell: object, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(
    cell: object, top: int = 55, start: int = 70, bottom: int = 55, end: int = 70
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _repeat_table_header(row: object) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _keep_row_together(row: object) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _configure_styles(doc: DocumentType) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.08
    for style_name, size, color, before, after in (
        ("Heading 1", 14, BLUE, 10, 5),
        ("Heading 2", 11.5, DARK_BLUE, 7, 3),
    ):
        style = doc.styles[style_name]
        style.font.name = LATIN_FONT
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def _configure_page(doc: DocumentType, blueprint: CourseBlueprint) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.48)
    section.bottom_margin = Inches(0.46)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.22)
    section.footer_distance = Inches(0.22)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(f"LessonCraft AI · 精简授课版 · {blueprint.course.title}")
    _set_run_font(run, size=7.8, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("全部教学事实由 Course Blueprint 统一派生")
    _set_run_font(run, size=7.8, color=MUTED)


def _add_title(doc: DocumentType, blueprint: CourseBlueprint) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(1)
    run = title.add_run(f"{blueprint.course.title}｜精简授课版教案")
    _set_run_font(run, size=20, bold=True, color=DARK_BLUE)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(7)
    run = subtitle.add_run(
        f"{blueprint.course.language}  ·  {blueprint.course.age_range}  ·  "
        f"{blueprint.course.duration_minutes} 分钟  ·  {blueprint.course.teaching_style}"
    )
    _set_run_font(run, size=10, color=MUTED)


def _cell_text(
    cell: object,
    text: str,
    *,
    size: float = 8.2,
    bold: bool = False,
    color: RGBColor = INK,
) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text or "—")
    _set_run_font(run, size=size, bold=bold, color=color)


def _add_meta_table(
    doc: DocumentType, rows: list[tuple[str, str, str, str]]
) -> None:
    table = doc.add_table(rows=0, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = [0.9, 4.05, 0.9, 4.05]
    for label_a, value_a, label_b, value_b in rows:
        cells = table.add_row().cells
        for index, cell in enumerate(cells):
            cell.width = Inches(widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)
        for index in (0, 2):
            _set_cell_fill(cells[index], TABLE_FILL)
        _cell_text(cells[0], label_a, bold=True, color=DARK_BLUE)
        _cell_text(cells[1], value_a)
        _cell_text(cells[2], label_b, bold=True, color=DARK_BLUE)
        _cell_text(cells[3], value_b)


def _add_inline_list(doc: DocumentType, label: str, items: list[str]) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(f"{label}：")
    _set_run_font(run, size=9.2, bold=True, color=DARK_BLUE)
    run = paragraph.add_run("；".join(items) if items else "无")
    _set_run_font(run, size=9.2, color=INK)


def _slide_pages(stage_id: str, slide_deck: SlideDeck) -> str:
    """Return human-readable page numbers derived from the current deck order."""

    pages = [
        index
        for index, slide in enumerate(slide_deck.slides, start=1)
        if stage_id in slide.source_step_ids
    ]
    if not pages:
        return "—"
    ranges: list[str] = []
    start = previous = pages[0]
    for page in pages[1:]:
        if page == previous + 1:
            previous = page
            continue
        ranges.append(str(start) if start == previous else f"{start}～{previous}")
        start = previous = page
    ranges.append(str(start) if start == previous else f"{start}～{previous}")
    return f"第 {'、'.join(ranges)} 页"


def _add_flow_table(
    doc: DocumentType, lesson_plan: LessonPlan, slide_deck: SlideDeck
) -> None:
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(4)
    run = note.add_run(
        "时间口径：教案时间包含投屏讲解、学生操作、教师巡视和环节过渡；"
        "PPT 备注时间仅表示逐页投屏讲解建议时间。"
    )
    _set_run_font(run, size=8.4, color=MUTED)
    headers = ["时间", "对应PPT页", "教学环节", "教师活动", "学生活动", "对应目标", "材料或代码"]
    widths = [0.55, 0.75, 1.25, 2.7, 2.3, 0.9, 1.45]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    header_row = table.rows[0]
    _repeat_table_header(header_row)
    for index, (cell, header) in enumerate(zip(header_row.cells, headers, strict=True)):
        cell.width = Inches(widths[index])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_fill(cell, TABLE_FILL)
        _set_cell_margins(cell)
        _cell_text(cell, header, size=8, bold=True, color=DARK_BLUE)
    for row_index, stage in enumerate(lesson_plan.stages):
        row = table.add_row()
        _keep_row_together(row)
        values = [
            f"{stage.duration.total_minutes} 分钟",
            _slide_pages(stage.step_id, slide_deck),
            stage.title,
            stage.teacher_activity,
            stage.student_activity,
            "、".join(stage.objective_ids),
            "、".join(stage.materials_or_code) or "—",
        ]
        for index, (cell, value) in enumerate(zip(row.cells, values, strict=True)):
            cell.width = Inches(widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)
            if row_index % 2:
                _set_cell_fill(cell, ALT_FILL)
            _cell_text(cell, value, size=8.1, bold=index == 2)


def _add_code_appendix(doc: DocumentType, blueprint: CourseBlueprint) -> None:
    if not blueprint.code_examples:
        return
    doc.add_heading("六、示例代码（每份仅出现一次）", level=1)
    for item in blueprint.code_examples:
        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(4)
        heading.paragraph_format.space_after = Pt(2)
        run = heading.add_run(f"{item.id} · {item.title}")
        _set_run_font(run, size=10, bold=True, color=DARK_BLUE)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        code_cell, note_cell = table.rows[0].cells
        code_cell.width = Inches(6.75)
        note_cell.width = Inches(3.15)
        _set_cell_fill(code_cell, "F3F5F7")
        for cell in (code_cell, note_cell):
            _set_cell_margins(cell, top=70, start=90, bottom=70, end=90)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        code_para = code_cell.paragraphs[0]
        code_para.paragraph_format.space_after = Pt(0)
        run = code_para.add_run(item.code)
        _set_run_font(
            run,
            name=CODE_FONT,
            east_asia=CJK_FONT,
            size=8.3,
            color=INK,
        )
        _cell_text(
            note_cell,
            f"用途：{item.explanation}\n目标：{'、'.join(item.objective_ids)}\n知识：{'、'.join(item.knowledge_ids)}",
            size=8.1,
        )


def _add_exercise_table(doc: DocumentType, blueprint: CourseBlueprint) -> None:
    """List observable practice tasks without duplicating answers or slide text."""

    if not blueprint.exercises:
        return
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    headers = ["练习 ID", "交付方式", "难度", "课堂练习或检查问题", "对应目标"]
    widths = [0.8, 1.15, 0.65, 6.0, 1.2]
    delivery_labels = {
        "in_class": "课堂检查",
        "student_assignment": "正式作业",
        "teacher_optional": "教师追加",
        "extension_challenge": "投屏挑战",
    }
    for index, (cell, header) in enumerate(zip(table.rows[0].cells, headers, strict=True)):
        cell.width = Inches(widths[index])
        _set_cell_fill(cell, TABLE_FILL)
        _set_cell_margins(cell)
        _cell_text(cell, header, size=8.1, bold=True, color=DARK_BLUE)
    _repeat_table_header(table.rows[0])
    for row_index, item in enumerate(blueprint.exercises):
        row = table.add_row()
        _keep_row_together(row)
        values = [
            item.id,
            delivery_labels[item.delivery_mode],
            item.difficulty,
            item.question,
            "、".join(item.objective_ids),
        ]
        for index, (cell, value) in enumerate(zip(row.cells, values, strict=True)):
            cell.width = Inches(widths[index])
            _set_cell_margins(cell)
            if row_index % 2:
                _set_cell_fill(cell, ALT_FILL)
            _cell_text(cell, value, size=8.2)


def export_lesson_plan_docx(
    blueprint: CourseBlueprint, lesson_plan: LessonPlan, slide_deck: SlideDeck
) -> bytes:
    """Export the default 3–5 page concise teaching plan as editable DOCX."""

    doc = Document()
    _configure_styles(doc)
    _configure_page(doc, blueprint)
    doc.core_properties.title = f"{blueprint.course.title}｜精简授课版教案"
    doc.core_properties.subject = "少儿编程精简授课教案"
    doc.core_properties.author = "LessonCraft AI"

    _add_title(doc, blueprint)
    doc.add_heading("一、课程基本信息", level=1)
    _add_meta_table(
        doc,
        [
            ("核心目标", blueprint.course.core_goal, "学生基础", blueprint.course.student_level),
            ("教学风格", blueprint.course.teaching_style, "PPT 页数", f"{len(slide_deck.slides)} 页"),
            ("补充要求", blueprint.course.additional_requirements or "无", "课程概览", lesson_plan.course_overview),
        ],
    )
    doc.add_heading("二、教学目标", level=1)
    _add_inline_list(doc, "可观察目标", lesson_plan.teaching_objectives)
    doc.add_heading("三、教学重点与难点", level=1)
    _add_inline_list(doc, "重点", lesson_plan.key_points)
    _add_inline_list(doc, "难点", lesson_plan.difficult_points)
    doc.add_heading("四、课前准备", level=1)
    _add_inline_list(doc, "准备清单", lesson_plan.preparation)
    doc.add_heading("五、课堂流程表", level=1)
    _add_flow_table(doc, lesson_plan, slide_deck)
    doc.add_page_break()
    _add_code_appendix(doc, blueprint)
    doc.add_page_break()
    doc.add_heading("七、课堂评价与课后任务", level=1)
    _add_inline_list(doc, "评价方式", lesson_plan.classroom_assessment)
    _add_exercise_table(doc, blueprint)
    _add_inline_list(doc, "课后任务", lesson_plan.homework)
    doc.add_heading("八、教师提醒", level=1)
    _add_inline_list(doc, "授课提醒", lesson_plan.teacher_reminders)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
